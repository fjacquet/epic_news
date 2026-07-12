import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from epic_news.utils.holiday_report import assemble_holiday_docx
from epic_news.utils.holiday_report.assemble import MAX_ITINERARY_DAYS


class StubLLM:
    """Returns a day-list for the skeleton call, Markdown for everything else."""

    def call(self, messages):
        user = messages[-1]["content"]
        if "date" in messages[0]["content"] and "stops" in messages[0]["content"]:
            return '[{"date":"J1","label":"Montreux","stops":["Montreux"]}]'
        return f"## Section\nContenu pour: {user[:20]}"


class ManyDaysStubLLM:
    """Returns a hallucinated over-long day-list for the skeleton call."""

    def __init__(self, num_days: int):
        self.num_days = num_days

    def call(self, messages):
        user = messages[-1]["content"]
        if "date" in messages[0]["content"] and "stops" in messages[0]["content"]:
            days = [{"date": f"J{i}", "label": f"Jour {i}", "stops": []} for i in range(1, self.num_days + 1)]
            return json.dumps(days)
        return f"## Section\nContenu pour: {user[:20]}"


def _crew_result():
    outs = ["destination research", "accommodation+dining", "itinerary", "budget"]
    return SimpleNamespace(tasks_output=[SimpleNamespace(raw=o) for o in outs])


def _inputs():
    return {
        "destination": "France",
        "duration": "2 semaines",
        "family": "3",
        "origin": "Montreux",
        "user_preferences_and_constraints": "nature",
    }


def test_assemble_builds_docx_with_sections_and_days(tmp_path: Path):
    out = tmp_path / "guide.docx"
    path = assemble_holiday_docx(_crew_result(), _inputs(), str(out), llm=StubLLM())
    assert path == str(out)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Introduction" in text
    assert "Itinéraire" in text  # at least one day section
    assert "Budget" in text


def test_assemble_caps_itinerary_days(tmp_path: Path):
    """A hallucinated 100-day skeleton must not fan out past MAX_ITINERARY_DAYS fragments."""
    out = tmp_path / "guide.docx"
    path = assemble_holiday_docx(_crew_result(), _inputs(), str(out), llm=ManyDaysStubLLM(100))
    assert path == str(out)
    day_headings = [
        p.text
        for p in Document(str(out)).paragraphs
        if p.style is not None and p.style.name == "Heading 1" and "Itinéraire — Jour" in p.text
    ]
    assert 0 < len(day_headings) <= MAX_ITINERARY_DAYS
