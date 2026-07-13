from pathlib import Path

from docx import Document

from epic_news.utils.docx_report import build_docx


def _all_text(path: str) -> str:
    """All visible text in a DOCX, including table cells (paragraphs miss those)."""
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_build_docx_writes_headings_and_body(tmp_path: Path):
    out = tmp_path / "guide.docx"
    fragments = [
        ("Introduction", "Bienvenue à **Montreux**."),
        ("Jour 1", "- Départ\n- Route"),
    ]
    result = build_docx(fragments, {"title": "Carnet", "date": "2026-07-16"}, str(out))

    assert result == str(out)
    assert out.exists() and out.stat().st_size > 0
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Introduction" in text
    assert "Jour 1" in text
    assert "Montreux" in text


def test_build_docx_survives_thematic_break_fences(tmp_path: Path):
    """LLM fragments use `---` separators; pandoc must not read them as YAML.

    Two `---` lines bracketing emphasized text (`*...*`) previously made pandoc's
    markdown reader parse the span as a YAML metadata block and die with exit 64
    (`while scanning an alias`). Body prose must never be treated as metadata.
    """
    out = tmp_path / "guide.docx"
    fragments = [
        (
            "Introduction",
            "Bienvenue.\n\n---\n*Note importante*\nTexte\n---\n\nFin.",
        ),
    ]
    result = build_docx(fragments, {"title": "Carnet", "date": "2026-07-16"}, str(out))

    assert result == str(out)
    assert out.exists() and out.stat().st_size > 0
    text = _all_text(str(out))
    assert "Note importante" in text
    assert "Fin." in text
